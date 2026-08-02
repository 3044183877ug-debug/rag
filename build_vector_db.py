"""
build_vector_db.py
------------------
YAML 驱动的税务政策向量化工具。

读取 data_source.yaml 作为"总索引"，遍历每个 policy entry，
定位 tax law documents/ 下的物理文件，加载全文后使用
RecursiveCharacterTextSplitter 切分，并将 YAML 中的全部元数据字段
注入每一个 Document 分块，最后存入 Chroma 向量数据库持久化。

支持的物理文件格式:
  - .txt   → 直接 UTF-8 读取
  - .pdf   → pymupdf (fitz) 逐页提取纯文本
  - .docx  → python-docx 提取文本 + 表格转 Markdown（保留二维结构）
  - .doc   → win32com COM 自动化中转 → 另存为 .docx → 复用 load_docx
  - 附件  → 通过 YAML attachments 字段绑定到主文件，拼接到正文后统一切分
"""

import os
import re
import shutil
import sys
from pathlib import Path

import yaml
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma

# ── 文件格式依赖（顶部导入，启动即检测，避免构建中途静默失败）──
try:
    import fitz  # pymupdf — PDF 加载
except ImportError:
    print("[FATAL] 缺少 pymupdf，请执行: pip install pymupdf")
    sys.exit(1)

try:
    from docx import Document as DocxDocument  # python-docx — .docx 加载
    from docx.oxml.ns import qn as docx_qn
except ImportError:
    print("[FATAL] 缺少 python-docx，请执行: pip install python-docx")
    sys.exit(1)

# win32com 仅 Windows 环境需要（用于 .doc 旧格式），缺包时不阻断启动，
# 仅在遇到 .doc 文件时才报错
_HAS_WIN32COM = False
try:
    import pythoncom
    import win32com.client
    _HAS_WIN32COM = True
except ImportError:
    pass

# ══════════════════════════════════════════════════════════════════
#  路径配置
# ══════════════════════════════════════════════════════════════════
BASE_DIR = Path(__file__).resolve().parent
YAML_PATH = BASE_DIR / "data_source.yaml"
DOCS_DIR = BASE_DIR / "tax law documents"
DB_DIR = BASE_DIR / "chroma_db"

# ══════════════════════════════════════════════════════════════════
#  切分配置
# ══════════════════════════════════════════════════════════════════
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 350
# ── 切分符设计原则 ──────────────────────────────────────────────
# 法律文本长句极多，按逗号/分号切分会导致"主谓分离"。
# 例如："契税的纳税义务发生时间，为签订合同的当日。" → 切在逗号处丢失主语。
#
# 优先级递减（每级都是在前级无法切出 ≤ CHUNK_SIZE 的片段时才触发）：
#   \n\n  → 段落边界（最安全）
#   \n    → 换行/表格行边界（优先于句号，防止表格行被截断）
#   。    → 句子边界（法律文本的天然语义单元）
#   ；    → 分句边界（法律/公文常见分隔符）
#   |     → 表格列边界（防止长表格行被硬截在列内容中间）
#   " "   → 空格（极端情况兜底）
#
# RecursiveCharacterTextSplitter 内置兜底：
#   所有 5 级 separators 均失效时，LangChain 自动回退到 chunk_size 字符级硬切，
#   确保任何情况下单个 chunk 绝对不会超过 chunk_size 个字符。
CHUNK_SEPARATORS = ["\n\n", "\n", "。", "；", "|", " "]

# ── _normalize_text 与 Splitter 的交互审计 ─────────────────────
#
# Q: _normalize_text 的 Step 1（缝合跨页断句）会吃掉 \\n，是否导致
#    Splitter 失去切分点？
# A: 不会。Step 1 仅移除"非完结标点后的 \\n"（如逗号后跨页），
#    句号/问号/分号后的 \\n 全部保留。且 Step 2-3 会主动注入新的
#    \\n\\n 到"第X章""第X条""一、"等公文标记前，增加切分点密度。
#
# Q: 如果所有 separators 都找不到（纯 URL、长数字串等），chunk 会超限吗？
# A: 不会。RecursiveCharacterTextSplitter 的递归算法保证：遍历完
#    separators 列表后若仍超限，则按 chunk_size 字符硬截断。这是
#    LangChain 源码级的兜底保证，无需额外代码。
# ══════════════════════════════════════════════════════════════════


# ══════════════════════════════════════════════════════════════════
#  文件加载器
# ══════════════════════════════════════════════════════════════════

def load_txt(file_path: Path) -> str:
    """加载 .txt 文件（UTF-8 编码）"""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def load_pdf(file_path: Path) -> str:
    """加载 .pdf 文件，逐页提取纯文本并重建段落结构。

    pymupdf 的 get_text() 按视觉布局输出，每一视觉行都以 \\n 结尾，
    导致法条句子被大量假换行截断（每 40~60 字一个 \\n）。

    本函数做两件事：
    1. 移除段落内部的假换行：单个 \\n（前后均非 \\n）→ 删除，使句子重新连贯
    2. 保留真正的段落边界：连续两个以上 \\n → 保持不变
    """
    doc = fitz.open(str(file_path))
    if doc.is_encrypted:
        # 尝试空密码
        doc.authenticate("")

    texts: list[str] = []
    for page in doc:
        text = page.get_text()
        if not text.strip():
            continue

        # ── 核心：段落重建 ──────────────────────────────────────
        # 将孤立的 \\n（前后都不是 \\n）视为 PDF 视觉换行，直接删除
        # 连续 \\n\\n 或更多才是真正的段落分界，原样保留
        text = re.sub(r'(?<!\n)\n(?!\n)', '', text)
        texts.append(text.strip())

    doc.close()
    return "\n\n".join(texts)


# ══════════════════════════════════════════════════════════════════
#  .doc / .docx 加载器（保留表格 2D 结构）
# ══════════════════════════════════════════════════════════════════

def _table_to_markdown(table) -> str:
    """将 python-docx Table 转为 Markdown 表格，保留二维行列结构。

    说明：
      - python-docx 对纵向合并单元格会在每行重复返回相同的 cell.text，
        因此无需额外的前向填充逻辑
      - 单元格内换行全部压缩为空格，确保 Markdown 单行不跨行
      - 行宽不一致时（如遇横向合并单元格），用空字符串补齐到最大列数
    """
    if not table.rows:
        return ""

    # 1. 提取各行数据
    rows_data: list[list[str]] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = cell.text.strip()
            text = " ".join(text.split())  # 压缩连续空白、换行 → 单空格
            cells.append(text)
        rows_data.append(cells)

    # 2. 确定列数（取最大行宽）
    ncols = max(len(r) for r in rows_data) if rows_data else 0
    if ncols == 0:
        return ""

    # 3. 补齐列宽（仅用空串补齐，不做前向填充）
    for row in rows_data:
        while len(row) < ncols:
            row.append("")

    # 4. 组装 Markdown 表格
    lines: list[str] = []
    lines.append("| " + " | ".join(rows_data[0]) + " |")
    lines.append("| " + " | ".join(["---"] * ncols) + " |")
    for row in rows_data[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def load_docx(file_path: Path) -> str:
    """加载 .docx 文件，段落与表格交替提取，表格以 Markdown 保留二维结构。

    核心挑战：python-docx 的 doc.paragraphs 和 doc.tables 是分离的两个列表，
    无法还原文档中段落/表格的真实穿插顺序。

    解法：遍历 doc.element.body 的 XML 子元素，按 w:p (段落) / w:tbl (表格)
    标签判断类型，保持原始文档顺序交替提取。
    """
    TAG_P = docx_qn("w:p")
    TAG_TBL = docx_qn("w:tbl")

    doc = DocxDocument(str(file_path))
    parts: list[str] = []

    tables = doc.tables
    paragraphs = doc.paragraphs
    p_idx = 0
    t_idx = 0

    for child in doc.element.body:
        if child.tag == TAG_P:
            if p_idx < len(paragraphs):
                text = paragraphs[p_idx].text.strip()
                if text:
                    parts.append(text)
                p_idx += 1
        elif child.tag == TAG_TBL:
            if t_idx < len(tables):
                md_table = _table_to_markdown(tables[t_idx])
                if md_table:
                    parts.append(md_table)
                t_idx += 1

    return "\n\n".join(parts)


def load_doc(file_path: Path) -> str:
    """加载旧版 .doc（二进制 OLE 格式）通过 Windows COM 自动化中转。

    python-docx 只能读取 Open XML 格式 (.docx)，无法直接处理旧式 .doc。
    本函数利用已安装的 Microsoft Word COM 对象：
      1. 静默打开 .doc
      2. SaveAs 为临时 .docx (FileFormat=16)
      3. 关闭并退出 Word
      4. 调用 load_docx() 提取内容
      5. 删除临时文件
    """
    if not _HAS_WIN32COM:
        raise ImportError("缺少 pywin32，无法处理 .doc 文件。请执行: pip install pywin32")

    import tempfile

    file_path = file_path.resolve()
    co_init = False
    word = None
    tmp_path = None

    try:
        pythoncom.CoInitialize()
        co_init = True

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False  # 抑制"文件正在使用"等弹窗

        doc = word.Documents.Open(str(file_path), ReadOnly=True)

        # 创建临时 .docx 文件
        tmp_fd, tmp_path_str = tempfile.mkstemp(suffix=".docx", prefix="_rag_")
        os.close(tmp_fd)
        tmp_path = Path(tmp_path_str)

        doc.SaveAs2(str(tmp_path), FileFormat=16)  # 16 = wdFormatXMLDocument
        doc.Close()
    finally:
        if word:
            try:
                word.Quit()
            except Exception:
                pass
        if co_init:
            pythoncom.CoUninitialize()

    # 从临时 .docx 提取内容
    try:
        result = load_docx(tmp_path)
    finally:
        if tmp_path and tmp_path.exists():
            tmp_path.unlink()

    return result


# 文件扩展名 → 加载器映射
LOADER_MAP = {
    ".txt": load_txt,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".doc": load_doc,
}


# ══════════════════════════════════════════════════════════════════
#  元数据注入核心逻辑
# ══════════════════════════════════════════════════════════════════

# YAML 中需要从 metadata 排除的字段（文件名只用于定位，不注入）
METADATA_EXCLUDE_FIELDS = {"file_name", "attachments"}

# ══════════════════════════════════════════════════════════════════
#  字段归一化：YAML 中同时存在两套命名惯例，统一映射为规范名
#
#  惯例 A（研发条例等）：policy_name / source（文号） / applicable_scope
#  惯例 B（其他条目）：source_name / source_number / topic
#  ↓ 归一化后 ↓
#  规范名：       source_name / source_number / topic
# ══════════════════════════════════════════════════════════════════
FIELD_NORMALIZE = {
    "policy_name":      "source_name",
    "source":           "source_number",   # 注意：YAML 中 source 实际是文号
    "applicable_scope": "topic",
}

# 向下兼容别名：metadata 中生成 "source" 指向 source_name
# （rag_agent.py format_docs() 仍会 fallback 到 source）
METADATA_ALIAS_MAP = {
    "source_name": "source",
}


def _parse_article_chapter(text: str) -> dict:
    """
    从 chunk 文本中提取"章/节/条"结构信息。

    支持三种中国税务文档格式：
      - 法律类：第X章、第X节、第X条（如增值税法、契税法、个税法）
      - 行政法规类：第X章、第X条（如个税专项扣除暂行办法）
      - 通知/公告类：一、二、三、… 编号段落（如财税〔2015〕119号）

    Returns:
        {"chapter": str, "article": str}  — 缺省时为空字符串
    """
    # 中文数字 → 数值映射（用于解析"十条"、"二十条"等）
    _CN_DIGIT = {
        "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
        "六": 6, "七": 7, "八": 8, "九": 9,
    }

    def _cn_to_arabic(cn: str) -> str:
        """将中文数字转为阿拉伯数字字符串，无法转换则返回原字符串。"""
        cn = cn.strip()
        if not cn:
            return cn
        # 单位数
        if cn in _CN_DIGIT:
            return str(_CN_DIGIT[cn])
        # 十、十一...十九
        if cn == "十":
            return "10"
        if cn.startswith("十") and len(cn) == 2:
            return str(10 + _CN_DIGIT.get(cn[1], 0))
        # 二十、三十...九十
        if cn.endswith("十") and len(cn) == 2:
            return str(_CN_DIGIT.get(cn[0], 0) * 10)
        # 二十一...九十九
        if "十" in cn and len(cn) >= 3:
            parts = cn.split("十", 1)
            tens = _CN_DIGIT.get(parts[0], 0) * 10
            ones = _CN_DIGIT.get(parts[1], 0) if parts[1] else 0
            return str(tens + ones)
        return cn

    result = {"chapter": "", "article": ""}

    # ── 1. 提取"第X章" ──────────────────────────────────────────
    chapter_m = re.search(r'第([一二三四五六七八九十百千]+)章', text)
    if chapter_m:
        result["chapter"] = f"第{chapter_m.group(1)}章"

    # ── 2. 提取"第X条"（可能有多条，取首尾组成范围）────────────
    article_matches = re.findall(r'第([一二三四五六七八九十百千]+)条', text)
    if article_matches:
        first = article_matches[0]
        last = article_matches[-1] if len(article_matches) > 1 else first
        if first == last:
            result["article"] = f"第{first}条"
        else:
            result["article"] = f"第{first}条—第{last}条"

    # ── 3. 通知/公告类：没有"第X条"，提取"一、二、"等编号 ────
    if not result["article"] and not result["chapter"]:
        section_m = re.search(r'[（\(]?([一二三四五六七八九十]+)[）\)]?\s*[、．.]', text)
        if section_m:
            result["article"] = f"{section_m.group(1)}、"

    return result


def _is_table_sep(line: str) -> bool:
    """判断是否为 Markdown 表格分隔行 (| --- | --- |)"""
    return bool(re.fullmatch(r'\|\s*[-:]+(\s*\|\s*[-:]+)*\s*\|', line.strip()))


def _flatten_pipe_tables(text: str) -> str:
    """检测管道表格并逐行展平为自包含完整短句。

    将 Markdown 管道表格的每一数据行转为携带完整标题、税种上下文的自包含
    短句，消除 RecursiveCharacterTextSplitter 跨行切分表格导致的检索盲区。

    处理策略（按复杂度分三级）：
      1. 简单行（无分号子项）→ 直接展平："【表名】：税目xxx，税率为xxx。"
      2. 复杂行（含分号层级子项）→ 解析层级，每个叶子子项独立展平一行
      3. 无法解析的复杂行 → 保留原管道格式但注入表名前缀
    """
    lines = text.split('\n')
    output: list[str] = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ── 检测管道表格起始 ──
        if stripped.startswith('|') and not _is_table_sep(stripped):
            # 收集连续管道行
            table_lines: list[str] = []
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # 找表名：前面最近的非空、非管道行
            title = ""
            for j in range(len(output) - 1, -1, -1):
                candidate = output[j].strip()
                if candidate and not candidate.startswith('|') and not candidate.startswith('【'):
                    title = candidate.rstrip('：:')
                    break

            # 展平表格
            flattened = _flatten_table_block(title, table_lines)
            output.extend(flattened)
        else:
            output.append(line)
            i += 1

    return '\n'.join(output)


def _flatten_table_block(title: str, table_lines: list[str]) -> list[str]:
    """展平单个表格块，返回替换用的行列表。

    管线：解析列头 → 跳过标题行+分隔行 → 逐行展平
    """
    if len(table_lines) < 3:
        # 表格太小（没有数据行），原样保留
        return table_lines

    result: list[str] = []

    # ── 1. 解析列头 ──
    header_line = table_lines[0]
    headers = [col.strip() for col in header_line.split('|')[1:-1]]

    # ── 2. 找到数据行起始位置（跳过分隔行 + 可能的重复表头） ──
    data_start = 1
    while data_start < len(table_lines):
        line = table_lines[data_start]
        if _is_table_sep(line):
            data_start += 1
            continue
        # 跳过与 header 内容重复的表头行（跨页表格常见）
        line_cols = [col.strip() for col in line.split('|')[1:-1]]
        if line_cols == headers:
            data_start += 1
            continue
        break

    # ── 3. 逐行展平 ──
    for idx in range(data_start, len(table_lines)):
        row_line = table_lines[idx]
        cols = [col.strip() for col in row_line.split('|')[1:-1]]
        if not cols or all(c == '' for c in cols):
            continue

        flattened_rows = _flatten_one_row(title, headers, cols)
        result.extend(flattened_rows)

    return result


def _flatten_one_row(title: str, headers: list[str], cols: list[str]) -> list[str]:
    """展平单行表格数据。

    Returns:
        展平后的文本行列表（简单行返回 1 行，复杂行可能返回多行）
    """
    # 确保 cols 与 headers 对齐（补齐缺列）
    while len(cols) < len(headers):
        cols.append("")

    # ── 合并多列为描述性文本 ──
    # 逻辑：第一列通常是"税目/污染物名"，后续列为"税率/税额/备注"
    item_desc = cols[0] if cols else ""
    value_parts = []
    for j in range(1, len(cols)):
        hdr = headers[j] if j < len(headers) else ""
        val = cols[j]
        if val:
            if hdr and hdr != "备 注":
                value_parts.append(f"{hdr}{val}")
            elif val:
                value_parts.append(val)

    value_str = "，".join(value_parts) if value_parts else ""

    # ── 情况 1：简单行（税目列不含分号层级） ──
    if '；' not in item_desc:
        # 清理中文/阿拉伯数字序号前缀
        clean_desc = re.sub(r'^[一二三四五六七八九十]+、', '', item_desc)
        clean_desc = re.sub(r'^\d+[\.、]\s*', '', clean_desc)
        sentence = _make_sentence(title, clean_desc, value_str)
        return [sentence]

    # ── 情况 2：复杂行（含分号层级子项） ──
    return _flatten_complex_row(title, headers, cols)


def _flatten_complex_row(title: str, headers: list[str], cols: list[str]) -> list[str]:
    """处理含分号层级子项的复杂表格行。

    子项格式如：
      税目: "一、烟；1.卷烟；（1）甲类卷烟；（2）乙类卷烟；2.雪茄烟；3.烟丝"
      税率: "；；45%加0.003元/支；30%加0.003元/支；25%；30%"

    策略：按层级收集叶子项（具体应税品目→税率），逐个展平。
    """
    item_desc = cols[0] if cols else ""
    # 收集后续列的原始值
    rate_cols = cols[1:] if len(cols) > 1 else []

    # ── 解析税目列为 tokens: [(level, text), ...] ──
    # 层级判定：数字+顿号=一级, 数字+点=二级, 带括号=三级叶子
    parts = [p.strip() for p in item_desc.split('；')]
    raw_tokens: list[dict] = []  # {text, level, is_leaf}
    for p in parts:
        if not p:
            continue
        # 判断层级
        if re.match(r'^[一二三四五六七八九十]+、', p):
            level = 1
        elif re.match(r'^\d+\.', p):
            level = 2
        elif re.match(r'^[（(]\d+[)）]', p):
            level = 3
        else:
            level = 2  # 默认
        raw_tokens.append({"text": p, "level": level})

    if not raw_tokens:
        # 降级：原行保留带表名
        return [_make_sentence(title, item_desc, "；".join(rate_cols))]

    # ── 为每个 token 标注叶子状态 ──
    # 叶子判定：level=3 或最后一个 token 或没有更深层级跟随
    for ti in range(len(raw_tokens)):
        t = raw_tokens[ti]
        if t["level"] == 3:
            t["is_leaf"] = True
        elif ti == len(raw_tokens) - 1:
            t["is_leaf"] = True
        elif ti + 1 < len(raw_tokens) and raw_tokens[ti + 1]["level"] > t["level"]:
            t["is_leaf"] = False
        else:
            t["is_leaf"] = True

    # ── 取叶子 tokens ──
    leaf_texts = [t for t in raw_tokens if t["is_leaf"]]

    # ── 解析税率列：按分号拆分子税率，跳过空的 ──
    all_rates: list[str] = []
    for rc in rate_cols:
        for sub in rc.split('；'):
            sub = sub.strip()
            if sub:
                all_rates.append(sub)

    # ── 匹配叶子到税率 ──
    # 税率数量应与叶子token数量一致，不一致时做 best-effort
    results: list[str] = []
    for li, leaf in enumerate(leaf_texts):
        rate = all_rates[li] if li < len(all_rates) else ""
        # 清理编号前缀
        clean_text = re.sub(r'^[（(]\d+[)）]\s*', '', leaf["text"])
        clean_text = re.sub(r'^\d+\.\s*', '', clean_text)
        clean_text = re.sub(r'^[一二三四五六七八九十]+、', '', clean_text)
        # 拼接父子路径：找上一级非叶子作为前缀
        prefix = ""
        for ti in range(len(raw_tokens)):
            if raw_tokens[ti] is leaf:
                break
            if not raw_tokens[ti].get("is_leaf"):
                ancestor = re.sub(r'^\d+[\.、]?\s*', '', raw_tokens[ti]["text"])
                ancestor = re.sub(r'^[一二三四五六七八九十]+、', '', ancestor)
                if ancestor and ancestor not in prefix:
                    prefix = prefix + ancestor + "-" if prefix else ancestor + "："
        full_desc = f"{prefix}{clean_text}"
        results.append(_make_sentence(title, full_desc, rate))

    return results if results else [_make_sentence(title, item_desc, "；".join(rate_cols))]


def _make_sentence(title: str, item_desc: str, value_str: str) -> str:
    """组装展平短句。"""
    if title:
        if value_str:
            return f"【{title}】：{item_desc}，{value_str}。"
        else:
            return f"【{title}】：{item_desc}。"
    else:
        if value_str:
            return f"{item_desc}，{value_str}。"
        else:
            return f"{item_desc}。"


def _normalize_text(text: str) -> str:
    """物理排版恢复 — 在切分前重建法律/财税文本的段落结构。

    解决的问题：
      a) PDF/OCR 跨页截断 — 一句话被硬生生切成两半，中间夹了一个换行符
      b) 文本"泥石流"   — 丢失全部段落换行，法条挤成一整块，splitter 无法切分

    处理管线（严格按顺序执行）：
      第一步  缝合跨页断句 — 非完结标点后的 \\n 一律视为异常截断，直接吃掉
      第二步  恢复法律排版 — 第X章 / 第X条 前注入段落边界
      第三步  恢复公文排版 — 大点（一、）中点（（一））小点（1.）前注入换行
      第四步  收尾压缩     — 连续 3+ 个 \\n 压缩为 \\n\\n

    设计目标：无论原始文本有多"脏"，经过本函数后 splitter 都能按语义边界
    干净切分，不再出现法条中截断、跨页残留、或段落粘连。
    """
    # ── 0. 行尾符归一化（Windows \\r\\n / 老 Mac \\r → Unix \\n）──
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = text.strip()

    # ════════════════════════════════════════════════════════════════
    # 第一步：缝合跨页断句
    #
    # 规则：扫描所有 \\n。如果 \\n 前面的字符不是中文完结标点
    #       （。！？；：、”、】），就认为这是一句被 PDF 跨页截断的话，
    #       直接吃掉这个 \\n，将前后文本无缝拼接。
    #
    # 例子：
    #   "纳税义务发生时间，\n为签订合同的当日。"  → 拼接为一句
    #   "本法自2026年1月1日起施行。\n第二条..."   → 句号后的 \\n 保留
    # ════════════════════════════════════════════════════════════════
    #
    # 采用捕获组方案（而非 lookbehind）：匹配"非完结标点 + \\n"，
    # 替换时只保留捕获的非完结标点字符，吃掉 \\n。
    #
    # 注意：[^...] 中不包含 \\n 本身，因此 \\n\\n 中第二个 \\n
    # 会被前一个 \\n "吃掉" → \\n\\n 降为 \\n，后续第四步统一处理。
    text = re.sub(
        r'([^。！？；：、”、】\n|])\n',
        r'\1',
        text,
    )

    # ════════════════════════════════════════════════════════════════
    # 第二步：恢复法律体排版（章、条）
    # ════════════════════════════════════════════════════════════════

    # 2a. 第X章 → \\n\\n 前导 + \\n 后随
    #     匹配 "第一章"、"第十二章"、"第一百二十章" 等
    text = re.sub(
        r'(第[一二三四五六七八九十百千]+章)',
        r'\n\n\1\n',
        text,
    )

    # 2b. 第X条 → \\n\\n 前导
    #     匹配 "第一条"、"第十五条"、"第一百二十条" 等
    text = re.sub(
        r'(第[一二三四五六七八九十百千]+条)',
        r'\n\n\1',
        text,
    )

    # ════════════════════════════════════════════════════════════════
    # 第三步：恢复财税公文体排版（大点 / 中点 / 小点）
    # ════════════════════════════════════════════════════════════════

    # 3a. 大点：中文数字 + 顿号 → \\n\\n 前导
    #     匹配 "一、" "二、" ... "十、" "十一、" ... "三十一、" 等
    text = re.sub(
        r'([一二三四五六七八九十]+、)',
        r'\n\n\1',
        text,
    )

    # 3b. 中点：带括号的中文数字 → \\n 前导
    #     兼容全角括号（一）（二）和半角括号 (一)(二)
    #     中文数字 1~3 个字符（一 到 九十九）
    text = re.sub(
        r'([（(][一二三四五六七八九十]{1,3}[）)])',
        r'\n\1',
        text,
    )

    # 3c. 小点：阿拉伯数字 + 点号 → \\n 前导
    #     兼容全角点号 1．2．和半角点号 1. 2.
    #     (?<!\d) 防止匹配 "3.14" 中的 "3." 部分
    #     (?!\d)  防止匹配 "3.14" 中的 ".1" 部分（点号后紧跟数字视为小数）
    text = re.sub(
        r'(?<!\d)(\d+)[．.](?!\d)',
        r'\n\1.',
        text,
    )

    # ════════════════════════════════════════════════════════════════
    # 第四步：收尾清理
    # ════════════════════════════════════════════════════════════════

    # 4a. 连续 3 个及以上 \\n → 压缩为 \\n\\n（恰好一个空行）
    text = re.sub(r'\n{3,}', '\n\n', text)

    # 4b. 去除首尾换行（步骤 2/3 可能在文本两端插入了前导/尾随 \\n\\n）
    return text.strip()


def _build_metadata(entry: dict) -> dict:
    """
    从 YAML entry 构建 metadata dict。

    处理流程：
    1. 字段归一化：将 YAML 中的两套命名惯例统一为规范名
       (policy_name → source_name, source → source_number, applicable_scope → topic)
    2. 排除 file_name（物理定位用，不注入语义元数据）
    3. list 类型转为逗号分隔字符串（Chroma 只接受标量）
    4. 生成 source 别名指向 source_name（兼容 rag_agent.py）
    5. 添加统一 date 字段（优先 date_effective，其次 date_published）
    6. 其他字段（如 status、doc_type、date_*、url）原样保留
    """
    metadata = {}

    for key, value in entry.items():
        if key in METADATA_EXCLUDE_FIELDS:
            continue

        # 字段归一化：将 YAML 中的异构命名映射为统一 key
        normalized_key = FIELD_NORMALIZE.get(key, key)

        # 将 list 类型转为逗号分隔字符串（Chroma 只接受标量）
        if isinstance(value, list):
            value = "，".join(str(v) for v in value)

        metadata[normalized_key] = value

    # 生成向下兼容的别名（如 source → source_name）
    for src_key, alias in METADATA_ALIAS_MAP.items():
        if src_key in metadata and alias not in metadata:
            metadata[alias] = metadata[src_key]

    # 生成统一 date 字段（优先生效日期，其次发布日期）
    if "date" not in metadata:
        metadata["date"] = metadata.get("date_effective") or metadata.get("date_published") or ""

    return metadata


def load_and_chunk(entry: dict) -> list[Document]:
    """
    加载单个 YAML entry 对应的物理文件，切分并注入元数据。

    步骤：
    1. 从 entry["file_name"] 获取文件名
    2. 在 DOCS_DIR 下查找物理文件
    3. 根据扩展名选择加载器读取全文
    4. 若 entry 含 attachments 字段，逐附件加载并拼接到主文本末尾
       （附件内容以【附件：文件名】标记分隔，与主文本统一切分）
    5. 用 RecursiveCharacterTextSplitter 切分
    6. 将 entry 中除 file_name/attachments 外的所有字段注入每个 chunk
    7. 为每个 chunk 解析章/节/条信息并注入 metadata
    """
    file_name = entry["file_name"]
    file_path = DOCS_DIR / file_name

    if not file_path.exists():
        print(f"  [WARN]  跳过（文件不存在）: {file_path}")
        return []

    ext = file_path.suffix.lower()
    loader = LOADER_MAP.get(ext)

    if loader is None:
        print(f"  [WARN]  跳过（不支持的格式 .{ext}）: {file_path}")
        return []

    # 1. 加载主文件全文
    print(f"  读取: {file_name}")
    raw_text = loader(file_path)

    # 2. 加载附件（如有），拼接到主文本末尾后统一切分
    attachment_files = entry.get("attachments", [])
    if attachment_files:
        for att_file in attachment_files:
            att_path = DOCS_DIR / att_file
            if not att_path.exists():
                print(f"    [WARN]  附件不存在: {att_file}")
                continue
            att_ext = att_path.suffix.lower()
            att_loader = LOADER_MAP.get(att_ext)
            if att_loader is None:
                print(f"    [WARN]  附件格式不支持 (.{att_ext}): {att_file}")
                continue
            print(f"    [ATT] 附件: {att_file}")
            try:
                att_text = att_loader(att_path)
                # 用明确标记分隔，方便 LLM 识别附件边界
                raw_text += f"\n\n【附件：{att_path.stem}】\n{att_text}"
            except Exception as e:
                print(f"    [ERROR] 附件加载失败: {att_file} — {e}")

    # 3. 表格展平（必须在归一化之前！_normalize_text 的 Step 3a 会把表格内的
    #    "一、" "二、" 当成公文大点注入 \\n\\n，炸碎管道表格行）
    #    → 先展平为纯文本短句，后续归一化不会误伤
    #    → 归一化 + 切分
    text = _flatten_pipe_tables(raw_text)
    text = _normalize_text(text)
    print(f"    提取完整文本，共 {len(text)} 个字符")

    if not text:
        print(f"  [WARN]  跳过（文件内容为空）: {file_path}")
        return []

    # 2. 切分
    splitter = RecursiveCharacterTextSplitter(
        separators=CHUNK_SEPARATORS,
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        is_separator_regex=False,
    )
    chunks = splitter.split_text(text)
    print(f"    切分为 {len(chunks)} 个文本块")

    # ── Chunk 大小分布验证（安全阀：确认切分器未失效）──────
    if chunks:
        _lens = [len(c) for c in chunks]
        _max_l = max(_lens)
        _min_l = min(_lens)
        _avg_l = sum(_lens) // len(_lens)
        print(f"    Chunk 统计: 最小 {_min_l}  最大 {_max_l}  平均 {_avg_l}  上限 {CHUNK_SIZE}")
        if _max_l > CHUNK_SIZE * 1.2:
            print(f"    ⚠️ 警告: 存在超大 chunk ({_max_l} 字符)，超过 chunk_size 的 120%")
        else:
            print(f"    [OK] 全部 chunk 均在安全上限内")

    # 3. 构建元数据（基于 YAML entry）
    base_metadata = _build_metadata(entry)

    # 4. 全局上下文注入：提取文件标题，拼接为 chunk 前缀
    #    解决"语义词汇鸿沟"问题 — 包含具体数字（如"1500元"）的 chunk
    #    因缺少税种关键词而在向量检索中被挤出 Top-K。前缀让每个 chunk
    #    都携带所属文件的完整标题，embedding 计算时包含文件级语义锚点。
    doc_title = base_metadata.get("source_name") or base_metadata.get("source") or "未知文件"
    context_prefix = f"【所属文件标题：{doc_title}】\n"

    # 5. 组装 Document 对象，每个 chunk 附加独立的 article/chapter
    documents = []
    for chunk in chunks:
        chunk_meta = base_metadata.copy()
        # 解析当前 chunk 的章/条结构
        struct = _parse_article_chapter(chunk)
        if struct["chapter"]:
            chunk_meta["chapter"] = struct["chapter"]
        if struct["article"]:
            chunk_meta["article"] = struct["article"]
        # 注入上下文前缀 + 正文
        documents.append(Document(page_content=context_prefix + chunk, metadata=chunk_meta))

    return documents


# ══════════════════════════════════════════════════════════════════
#  主流程
# ══════════════════════════════════════════════════════════════════

def main():
    # ── 1. 加载 YAML ──────────────────────────────────────────────
    if not YAML_PATH.exists():
        print(f"[ERROR] 未找到 data_source.yaml: {YAML_PATH}")
        return

    with open(YAML_PATH, "r", encoding="utf-8") as f:
        entries = yaml.safe_load(f)

    if not entries:
        print("[ERROR] data_source.yaml 为空或格式错误")
        return

    print(f"[YAML] 加载 YAML 索引: {len(entries)} 个政策条目\n")

    # ── 2. 遍历 YAML entries，加载 + 切分 + 注入元数据 ──────────
    all_documents: list[Document] = []

    for i, entry in enumerate(entries, 1):
        source_name = entry.get("source_name", entry.get("file_name", "?"))
        print(f"[{i}/{len(entries)}] {source_name}")

        try:
            docs = load_and_chunk(entry)
            all_documents.extend(docs)
            print(f"    已生成 {len(docs)} 个 Document（累计: {len(all_documents)}）\n")
        except Exception as e:
            print(f"    [ERROR] 处理失败: {e}\n")
            continue

    if not all_documents:
        print("[ERROR] 没有生成任何文档片段，请检查文件路径和格式")
        return

    # ── 3. 打印预览 ──────────────────────────────────────────────
    print(f"\n{'='*60}")
    print(f"总计 {len(all_documents)} 个文档片段\n")

    # 统计各来源分布
    source_counts: dict[str, int] = {}
    for doc in all_documents:
        src = doc.metadata.get("source_name", "未知")
        source_counts[src] = source_counts.get(src, 0) + 1

    print("来源分布:")
    for src, count in sorted(source_counts.items(), key=lambda x: -x[1]):
        print(f"  {count:>4}  | {src}")

    # 打印前 5 条样本
    print(f"\n前 5 条片段预览:\n")
    for i, doc in enumerate(all_documents[:5], 1):
        meta = doc.metadata
        src = meta.get("source_name", "?")[:30]
        num = meta.get("source_number", "")
        tp = meta.get("doc_type", "")
        preview = doc.page_content[:100].replace("\n", " ").replace("\r", " ")
        # 移除不可打印字符
        import unicodedata
        preview = "".join(c for c in preview if c.isprintable() or c in " ")
        print(f"  [{i}] {src}")
        if num:
            print(f"      文号: {num}")
        print(f"      类型: {tp}  |  生效: {meta.get('date_effective', '')}")
        topics = meta.get("topic", "")
        if topics:
            print(f"      主题: {topics}")
        print(f"      内容: {preview}...")
        print()

    # ── 4. 加载嵌入模型 ──────────────────────────────────────────
    print("正在加载 Embedding 模型（BAAI/bge-small-zh-v1.5）...")
    embedding_model = HuggingFaceEmbeddings(
        model_name="BAAI/bge-small-zh-v1.5",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )
    print("Embedding 模型加载完成\n")

    # ── 5. 存入 Chroma ───────────────────────────────────────────
    print(f"正在构建向量数据库，存储路径: {DB_DIR}")

    if DB_DIR.exists():
        shutil.rmtree(DB_DIR)
        print("  已清除旧的向量数据库")

    vectordb = Chroma.from_documents(
        documents=all_documents,
        embedding=embedding_model,
        persist_directory=str(DB_DIR),
        collection_name="tax_knowledge",
    )

    count = vectordb._collection.count()
    print(f"\n[OK] 向量数据库构建完成！")
    print(f"   集合名称: tax_knowledge")
    print(f"   向量总数: {count}")
    print(f"   持久化路径: {DB_DIR}")


if __name__ == "__main__":
    main()
